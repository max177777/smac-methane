"""
Builds data/rag_corpus.json — the chunked, tagged text corpus behind the
local RAG in utils/rag.py — from the internal methane action-plan library.

Two-pass pipeline:
  1. Extract text from every .docx/.pdf under RAW_SOURCE_DIR (+ a few
     root-level reference docs), tagging each with jurisdiction + tier.
     Flags likely-scanned PDFs (near-zero extractable text) instead of
     silently indexing garbage.
  2. Chunk each extracted document into ~180-word overlapping pieces, guess
     a sector tag via keyword matching, and default which chat "output
     types" (data/trend/policy/pathway/method) each chunk suits based on its
     tier.

RAW SOURCE FILES ARE NOT COMMITTED to this repo (the source zips run well
over 100MB — too large for a lean git-deployed app). Only the derived
data/rag_corpus.json (a few MB of chunked text) is committed. To rebuild
after a new document drop:

  1. Unzip the new source material into a local RAW_SOURCE_DIR (see below).
  2. Update JURISDICTION_MAP / ROOT_DOCS if new folders or root docs were added.
  3. Run: python scripts/build_rag_index.py
  4. Commit the updated data/rag_corpus.json.

Requires: pip install python-docx pdfplumber
"""

import json
import math
import os
import re
import signal
from collections import Counter
from pathlib import Path

import docx
import pdfplumber

# ---- Adjust these two paths for wherever the raw source material lives ----
RAW_SOURCE_DIR = Path("raw_sources")          # unzipped document collections
ROOT_DOCS_DIR = Path("raw_sources_root")       # standalone reference docx files

OUT_PATH = Path(__file__).parent.parent / "data" / "rag_corpus.json"

CHUNK_WORDS = 180
OVERLAP_WORDS = 30
PDF_PER_FILE_TIMEOUT_SEC = 45

# folder-name substring -> (iso3, location or None, tier)
# tier is one of: smac_member | national | solution_bank | template | reference_catalog | other
JURISDICTION_MAP = [
    ("Emma_s collection", "MEX", "Jalisco", "smac_member"),
    ("Maryland Resources", "USA", "Maryland", "smac_member"),
    ("Mexico(National) Resources", "MEX", None, "national"),
    ("Methane Action Plans and Template/Maryland Methane Action Plan drafts", "USA", "Maryland", "smac_member"),
    ("Methane Action Plans and Template/Solution Bank Archive", None, None, "solution_bank"),
    ("Methane Action Plans and Template/Methane Action Plan Template", None, None, "template"),
]

# (filename in ROOT_DOCS_DIR, iso3, location, tier)
ROOT_DOCS = [
    ("Current_Methane_Action_Plans.docx", None, None, "reference_catalog"),
    ("Subnational_Methane_Action_Plan_Reference_Library.docx", None, None, "reference_catalog"),
    ("Methane_Solutions_draft_1_.docx", None, None, "solution_bank"),
]

SECTOR_KEYWORDS = {
    "Agriculture": ["livestock", "enteric", "manure", "rice cultivation", "dairy", "cattle",
                    "feed additive", "fertilizer", "agricultur"],
    "Waste": ["landfill", "wastewater", "solid waste", "msw", "composting", "organic waste",
              "waste diversion", "anaerobic digest"],
    "Fossil Fuel Extraction & Mining": ["oil and gas", "oil & gas", "flaring", "venting", "ldar",
                                        "coal mine", "pneumatic", "wellhead", "fugitive emission",
                                        "upstream", "well plugging", "orphaned well"],
    "Forestry & Land Use": ["forest", "land use", "deforestation", "wetland", "peatland", "redd+"],
    "Manufacturing & Industry": ["cement", "steel", "chemical", "manufactur", "industrial process"],
    "Power & Heat": ["electricity generation", "power plant", "thermal power", "grid"],
    "Transportation": ["transportation", "vehicle fleet", "transit", "aviation", "shipping"],
    "Buildings (Onsite Fuel Use)": ["building", "residential fuel", "heating", "onsite fuel"],
}

TIER_OUTPUT_DEFAULTS = {
    "smac_member": ["policy", "pathway", "method", "data"],
    "national": ["policy", "pathway"],
    "solution_bank": ["pathway"],
    "template": ["method", "pathway"],
    "reference_catalog": ["policy"],
    "other": ["policy"],
}


# ============== PASS 1: EXTRACT ==============

def tag_for_path(rel_path: str):
    for prefix, iso, loc, tier in JURISDICTION_MAP:
        if prefix in rel_path:
            return iso, loc, tier
    return None, None, "other"


def extract_docx(path) -> str:
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


class _TimeoutErr(Exception):
    pass


def _alarm_handler(signum, frame):
    raise _TimeoutErr()


def extract_pdf(path, max_pages=400, per_file_timeout=PDF_PER_FILE_TIMEOUT_SEC):
    """Returns (text, status). status is one of OK / SCANNED_OR_EMPTY /
    TIMEOUT_PARTIAL / ERROR: <msg>. SCANNED_OR_EMPTY means the average
    characters-per-page was too low to be real extracted text — almost
    certainly a scanned image PDF that would need OCR, not indexed here."""
    signal.signal(signal.SIGALRM, _alarm_handler)
    signal.alarm(per_file_timeout)
    text_parts, n_pages, n_chars = [], 0, 0
    try:
        with pdfplumber.open(path) as pdf:
            n_pages = len(pdf.pages)
            for page in pdf.pages[:max_pages]:
                t = page.extract_text() or ""
                text_parts.append(t)
                n_chars += len(t)
        signal.alarm(0)
    except _TimeoutErr:
        signal.alarm(0)
        return "\n".join(text_parts), "TIMEOUT_PARTIAL"
    except Exception as e:
        signal.alarm(0)
        return "", f"ERROR: {e}"
    full = "\n".join(text_parts)
    chars_per_page = n_chars / max(n_pages, 1)
    return full, ("SCANNED_OR_EMPTY" if chars_per_page < 30 else "OK")


def extract_all():
    """Yields dicts: rel_path, filename, iso3, location, tier, status, text."""
    for fname, iso, loc, tier in ROOT_DOCS:
        path = ROOT_DOCS_DIR / fname
        if not path.exists():
            print(f"  (skip, not found) {path}")
            continue
        text = extract_docx(path)
        yield {"rel_path": f"__root__/{fname}", "filename": fname, "iso3": iso,
               "location": loc, "tier": tier, "status": "OK" if text.strip() else "EMPTY",
               "text": text}

    if not RAW_SOURCE_DIR.exists():
        print(f"  (skip) RAW_SOURCE_DIR not found: {RAW_SOURCE_DIR}")
        return

    for root, _, files in os.walk(RAW_SOURCE_DIR):
        for fname in files:
            path = Path(root) / fname
            rel = str(path.relative_to(RAW_SOURCE_DIR))
            ext = fname.rsplit(".", 1)[-1].lower()
            iso, loc, tier = tag_for_path(rel)
            if ext == "docx":
                try:
                    text = extract_docx(path)
                    status = "OK" if text.strip() else "EMPTY"
                except Exception as e:
                    text, status = "", f"ERROR: {e}"
            elif ext == "pdf":
                text, status = extract_pdf(path)
            else:
                continue
            print(f"  [{status:>16}] {rel} ({len(text)} chars)")
            yield {"rel_path": rel, "filename": fname, "iso3": iso, "location": loc,
                   "tier": tier, "status": status, "text": text}


# ============== PASS 2: CHUNK + TAG ==============

def guess_sector(text_lower: str):
    scores = {}
    for sector, kws in SECTOR_KEYWORDS.items():
        c = sum(text_lower.count(kw) for kw in kws)
        if c:
            scores[sector] = c
    return max(scores, key=scores.get) if scores else None


def chunk_words(words, size, overlap):
    i = 0
    while i < len(words):
        yield words[i:i + size]
        if i + size >= len(words):
            break
        i += size - overlap


def build_chunks(docs):
    chunks, cid = [], 0
    for doc in docs:
        if doc["status"] != "OK" or not doc["text"].strip():
            continue
        text = re.sub(r"\s+", " ", doc["text"]).strip()
        words = text.split(" ")
        for piece in chunk_words(words, CHUNK_WORDS, OVERLAP_WORDS):
            piece_text = " ".join(piece).strip()
            if len(piece_text) < 80:
                continue
            chunks.append({
                "id": cid,
                "text": piece_text,
                "source_file": doc["filename"],
                "source_path": doc["rel_path"],
                "iso3": doc["iso3"],
                "location": doc["location"],
                "tier": doc["tier"],
                "sector": guess_sector(piece_text.lower()),
                "output_types": TIER_OUTPUT_DEFAULTS.get(doc["tier"], ["policy"]),
            })
            cid += 1
    return chunks


def main():
    print("Pass 1/2: extracting text...")
    docs = list(extract_all())
    ok = [d for d in docs if d["status"] == "OK"]
    print(f"  {len(ok)}/{len(docs)} documents extracted successfully.")

    print("Pass 2/2: chunking + tagging...")
    chunks = build_chunks(docs)
    print(f"  {len(chunks)} chunks.")
    print("  by tier:", Counter(c["tier"] for c in chunks))
    print("  by location:", Counter(c["location"] for c in chunks))
    print("  by sector:", Counter(c["sector"] for c in chunks))

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(OUT_PATH, "w", encoding="utf-8") as f:
        json.dump(chunks, f)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
